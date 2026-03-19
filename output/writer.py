"""
Output Writer
-------------
Fills the SSC1 Risk Review table (document_for_format.docx template) with:
  Col 1: S.No
  Col 2: Clause name
  Col 3: Clause Reference + extracted clause text
  Col 4: Risk Involved
  Col 5: R&Q Remarks (auto-filled)

Uses python-docx to write into the pre-formatted template.
"""

import copy
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dataclasses import dataclass
from typing import List, Optional


@dataclass
class TableRow:
    sno: str
    clause_name: str
    clause_reference: str        # "Clause 4.1, Page 12"
    original_clause_text: str    # verbatim extracted text
    risk_involved: str           # risk description
    risk_level: str              # HIGH / MEDIUM / ACCEPTABLE / NEEDS_REVIEW
    rq_remarks: str              # auto-filled R&Q remark
    needs_exception: bool = False
    needs_eqcr: bool = False
    deviation_suggested: str = ""


# Risk level → colour mapping (for the Risk Involved cell)
RISK_COLORS = {
    "HIGH":         RGBColor(0xFF, 0xCC, 0xCC),   # light red
    "MEDIUM":       RGBColor(0xFF, 0xF0, 0xCC),   # light amber
    "ACCEPTABLE":   RGBColor(0xCC, 0xFF, 0xCC),   # light green
    "LOW":          RGBColor(0xCC, 0xFF, 0xCC),   # light green
    "NEEDS_REVIEW": RGBColor(0xCC, 0xCC, 0xFF),   # light blue
}

RISK_LABELS = {
    "HIGH":         "🔴 HIGH RISK",
    "MEDIUM":       "🟡 MEDIUM RISK",
    "ACCEPTABLE":   "🟢 ACCEPTABLE",
    "LOW":          "🟢 LOW RISK",
    "NEEDS_REVIEW": "🔵 NEEDS REVIEW",
}


def _set_cell_color(cell, rgb: RGBColor):
    """Apply background shading to a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def _add_cell_text(cell, text: str, bold: bool = False, size_pt: int = 9, color: Optional[RGBColor] = None):
    """Set cell text with formatting. Clears existing content first."""
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_multiline_cell(cell, lines: List[str], bold_first: bool = False, size_pt: int = 9):
    """Add multiple paragraphs to a cell."""
    first = True
    for line in lines:
        if not line.strip():
            continue
        if first:
            para = cell.paragraphs[0]
            para.clear()
            first = False
        else:
            para = cell.add_paragraph()

        run = para.add_run(line)
        run.font.size = Pt(size_pt)
        if bold_first and para == cell.paragraphs[0]:
            run.font.bold = True


def fill_ssc1_table(
    rows: List[TableRow],
    template_path: str,
    output_path: str,
    rfp_name: str = "",
) -> str:
    """
    Fill the SSC1 table template with extracted + risk-evaluated data.

    Args:
        rows: List of TableRow objects (one per clause)
        template_path: Path to document_for_format.docx
        output_path: Where to save the filled document
        rfp_name: Name of the RFP being reviewed (for header)

    Returns:
        output_path
    """
    # Copy template so we don't modify the original
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    if not doc.tables:
        raise ValueError("Template has no tables. Expected SSC1 table format.")

    table = doc.tables[0]

    # Find the first data row (skip header row)
    # Template already has header row + some example rows;
    # we'll clear example rows and write fresh data
    header_rows_count = 1  # adjust if template has more header rows

    # Remove existing data rows (keep only header)
    while len(table.rows) > header_rows_count:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    # Get column widths from header row for reference
    header_row = table.rows[0]
    col_count = len(header_row.cells)

    for idx, row_data in enumerate(rows, start=1):
        # Add a new row by copying the header row's structure
        new_row = copy.deepcopy(table.rows[0]._tr)
        table._tbl.append(new_row)
        row = table.rows[-1]

        # Ensure we have enough cells
        cells = row.cells
        if len(cells) < 5:
            continue

        # ── Col 0: S.No ─────────────────────────────────────────────────────
        _add_cell_text(cells[0], str(idx), bold=True, size_pt=9)

        # ── Col 1: Clause Name ───────────────────────────────────────────────
        clause_label = row_data.clause_name
        if row_data.needs_exception:
            clause_label += "\n[Exception Approval Required]"
        if row_data.needs_eqcr:
            clause_label += "\n[EQCR Applicable]"

        _add_cell_text(cells[1], clause_label, bold=True, size_pt=9)

        # ── Col 2: Clause Reference + Original Clause Text ───────────────────
        ref_text = row_data.clause_reference or "Not Found"
        clause_body = row_data.original_clause_text or "Clause not identified in document."
        # Truncate very long extractions for readability (keep first 800 chars)
        if len(clause_body) > 800:
            clause_body = clause_body[:797] + "..."

        _add_multiline_cell(
            cells[2],
            [f"Ref: {ref_text}", "", clause_body],
            bold_first=True,
            size_pt=8,
        )

        # ── Col 3: Risk Involved ─────────────────────────────────────────────
        risk_label = RISK_LABELS.get(row_data.risk_level, row_data.risk_level)
        risk_detail = row_data.risk_involved or ""
        _add_multiline_cell(
            cells[3],
            [risk_label, "", risk_detail],
            bold_first=True,
            size_pt=8,
        )

        # Colour the risk cell
        risk_color = RISK_COLORS.get(row_data.risk_level)
        if risk_color:
            _set_cell_color(cells[3], risk_color)

        # ── Col 4: R&Q Remarks ───────────────────────────────────────────────
        remarks_text = row_data.rq_remarks or "No auto-remarks. Manual review recommended."
        if row_data.deviation_suggested:
            remarks_text += f"\n\nSuggested Deviation Language:\n{row_data.deviation_suggested}"

        _add_cell_text(cells[4], remarks_text, size_pt=8)

    # Add a title/header note before the table if rfp_name provided
    if rfp_name:
        # Insert a paragraph before the table
        from docx.oxml import OxmlElement
        tbl_element = table._tbl
        p = OxmlElement("w:p")
        tbl_element.addprevious(p)

        new_para_idx = 0
        for i, para in enumerate(doc.paragraphs):
            if para._p == p:
                new_para_idx = i
                break

        doc.paragraphs[new_para_idx].text = f"SSC1 Risk & Quality Review — {rfp_name}"
        run = doc.paragraphs[new_para_idx].runs[0]
        run.bold = True
        run.font.size = Pt(12)

    doc.save(output_path)
    print(f"[OutputWriter] Saved SSC1 document: {output_path}")
    return output_path


# ──────────────────────────────────────────────────────────────────────────────
# Convert pipeline results → TableRow objects
# ──────────────────────────────────────────────────────────────────────────────

def build_table_rows(pipeline_results: dict) -> List[TableRow]:
    """
    Convert the output of the full pipeline (extractor + risk engine) into
    TableRow objects ready for the DOCX writer.

    pipeline_results format:
    {
        "liability": {
            "extracted": {...},
            "risk": RiskResult,
            ...
        },
        ...
    }
    """
    rows = []
    # Define display order
    CLAUSE_ORDER = [
        ("liability",    "Limitation of Liability"),
        ("insurance",    "Insurance Clause"),
        ("scope",        "Scope of Work"),
        ("payment",      "Payment Terms"),
        ("deliverables", "Deliverables"),
        ("personnel",    "Replacement/Substitution of Personnel/Key Resources"),
        ("ld",           "Liquidated Damages"),
        ("penalties",    "Penalties"),
        ("termination",  "Termination Rights"),
        ("eligibility",  "Eligibility Clause"),
    ]

    for clause_key, display_name in CLAUSE_ORDER:
        result = pipeline_results.get(clause_key, {})
        extracted = result.get("extracted", {})
        risk = result.get("risk")

        # Build clause reference string
        clause_ref = extracted.get("clause_reference", "")
        page_no = extracted.get("page_no", "")
        ref_str = ""
        if clause_ref:
            ref_str = str(clause_ref)
        if page_no:
            ref_str += f" (Page {page_no})" if ref_str else f"Page {page_no}"
        if not ref_str:
            ref_str = "Not identified"

        # Get original clause text
        clause_text = extracted.get("clause_text") or ""
        if clause_key == "scope":
            clause_text = extracted.get("summary") or clause_text

        # Risk fields
        if risk:
            risk_level = risk.risk_level
            risk_desc = risk.risk_description
            auto_remark = risk.auto_remark
            needs_exception = risk.needs_exception_approval
            needs_eqcr = risk.needs_eqcr
            deviation = risk.deviation_suggested
        else:
            risk_level = "NEEDS_REVIEW"
            risk_desc = "Risk evaluation not available."
            auto_remark = ""
            needs_exception = False
            needs_eqcr = False
            deviation = ""

        rows.append(TableRow(
            sno="",
            clause_name=display_name,
            clause_reference=ref_str,
            original_clause_text=clause_text,
            risk_involved=risk_desc,
            risk_level=risk_level,
            rq_remarks=auto_remark,
            needs_exception=needs_exception,
            needs_eqcr=needs_eqcr,
            deviation_suggested=deviation,
        ))

    return rows
