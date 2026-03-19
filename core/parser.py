"""
Document Parser
---------------
Extracts text from PDF or DOCX with page numbers and section headings.
Produces semantically meaningful chunks (by heading/section), NOT fixed-size chunks.
This is critical for RFPs which have natural clause structure.
"""

import re
import fitz  # pymupdf
from pathlib import Path
from docx import Document
from dataclasses import dataclass, field
from typing import List, Optional


@dataclass
class Chunk:
    """A single semantic chunk from the RFP document."""
    text: str
    page_no: int               # Page number in original doc
    section_heading: str       # Nearest heading (e.g. "4.3 Scope of Work")
    clause_ref: str            # E.g. "Clause 4.3" or "Section 5"
    doc_name: str              # Filename
    chunk_id: str              # Unique ID: {doc_name}_{page}_{seq}


def _extract_clause_ref(heading: str) -> str:
    """Extract a clean clause reference from a heading string."""
    if not heading:
        return ""
    # Match patterns like "4.3", "Clause 7", "Section 5.2.1", "Article 3"
    patterns = [
        r"(clause\s+[\d.]+)",
        r"(section\s+[\d.]+)",
        r"(article\s+[\d.]+)",
        r"(para(?:graph)?\s+[\d.]+)",
        r"([\d]+\.[\d]+(?:\.[\d]+)*)",
    ]
    for p in patterns:
        m = re.search(p, heading, re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return heading[:60]  # fallback to truncated heading


def _is_heading(text: str, font_size: Optional[float] = None, is_bold: Optional[bool] = None) -> bool:
    """Heuristic: is this line a section heading?"""
    text = text.strip()
    if not text or len(text) > 200:
        return False

    # Font-based detection (PDF)
    if font_size and font_size > 13:
        return True
    if is_bold and len(text) < 120:
        return True

    # Pattern-based detection
    heading_patterns = [
        r"^\d+\.\s+[A-Z]",               # "4. Scope"
        r"^\d+\.\d+\s+[A-Z]",            # "4.1 Limitation"
        r"^\d+\.\d+\.\d+\s+[A-Z]",       # "4.1.1 Sub-clause"
        r"^(Clause|Section|Article|Part|Schedule)\s+\d+",
        r"^[A-Z][A-Z\s]{5,50}$",          # ALL CAPS HEADINGS
        r"^(SCOPE|PAYMENT|LIABILITY|INSURANCE|TERMINATION|PENALTY|LIQUIDATED)",
    ]
    for p in heading_patterns:
        if re.match(p, text):
            return True
    return False


# ──────────────────────────────────────────────────────────────────────────────
# PDF Parser
# ──────────────────────────────────────────────────────────────────────────────

def parse_pdf(path: str) -> List[Chunk]:
    """
    Parse a PDF into semantic chunks grouped by section heading.
    Uses font size / bold detection for heading identification.
    """
    doc = fitz.open(path)
    doc_name = Path(path).name
    chunks: List[Chunk] = []

    current_heading = "Preamble"
    current_text_lines: List[str] = []
    current_page = 1
    seq = 0

    def flush_chunk():
        nonlocal seq
        text = "\n".join(current_text_lines).strip()
        if len(text) > 50:  # skip tiny fragments
            chunk = Chunk(
                text=text,
                page_no=current_page,
                section_heading=current_heading,
                clause_ref=_extract_clause_ref(current_heading),
                doc_name=doc_name,
                chunk_id=f"{doc_name}_{current_page}_{seq}",
            )
            chunks.append(chunk)
            seq += 1

    for page_num, page in enumerate(doc, start=1):
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] != 0:  # skip images
                continue
            for line in block["lines"]:
                line_text = " ".join(span["text"] for span in line["spans"]).strip()
                if not line_text:
                    continue

                # Detect heading from font properties
                max_size = max((s["size"] for s in line["spans"]), default=11)
                is_bold_line = any(s["flags"] & 2**4 for s in line["spans"])  # bit 4 = bold

                if _is_heading(line_text, font_size=max_size, is_bold=is_bold_line):
                    flush_chunk()
                    current_heading = line_text
                    current_text_lines = [line_text]
                    current_page = page_num
                else:
                    current_text_lines.append(line_text)

    flush_chunk()  # flush last section
    doc.close()
    return chunks


# ──────────────────────────────────────────────────────────────────────────────
# DOCX Parser
# ──────────────────────────────────────────────────────────────────────────────

def parse_docx(path: str) -> List[Chunk]:
    """
    Parse a DOCX into semantic chunks grouped by Heading-styled paragraphs.
    Falls back to pattern-based heading detection for unstyles docs.
    """
    doc = Document(path)
    doc_name = Path(path).name
    chunks: List[Chunk] = []

    current_heading = "Preamble"
    current_text_lines: List[str] = []
    current_page = 1  # DOCX doesn't expose page numbers easily; track roughly
    seq = 0
    para_count = 0

    # Rough page estimate: ~40 paragraphs per page in typical RFPs
    PARAS_PER_PAGE = 40

    def flush_chunk():
        nonlocal seq
        text = "\n".join(current_text_lines).strip()
        if len(text) > 50:
            chunk = Chunk(
                text=text,
                page_no=current_page,
                section_heading=current_heading,
                clause_ref=_extract_clause_ref(current_heading),
                doc_name=doc_name,
                chunk_id=f"{doc_name}_{current_page}_{seq}",
            )
            chunks.append(chunk)
            seq += 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        para_count += 1
        current_page = (para_count // PARAS_PER_PAGE) + 1

        # Check if paragraph uses a Heading style
        is_styled_heading = para.style.name.startswith("Heading")

        # Also check by text pattern
        is_pattern_heading = _is_heading(text)

        if is_styled_heading or is_pattern_heading:
            flush_chunk()
            current_heading = text
            current_text_lines = [text]
        else:
            current_text_lines.append(text)

    flush_chunk()
    return chunks


# ──────────────────────────────────────────────────────────────────────────────
# Also parse tables (important for PQ marking schemes!)
# ──────────────────────────────────────────────────────────────────────────────

def parse_docx_tables(path: str) -> List[Chunk]:
    """
    Extract tables from DOCX as text chunks.
    Many RFPs have marking schemes and eligibility criteria in tables.
    """
    doc = Document(path)
    doc_name = Path(path).name
    chunks: List[Chunk] = []
    seq = 0

    for table_idx, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows_text.append(" | ".join(cells))

        table_text = "\n".join(rows_text).strip()
        if len(table_text) > 30:
            chunks.append(Chunk(
                text=f"[TABLE {table_idx + 1}]\n{table_text}",
                page_no=0,
                section_heading=f"Table {table_idx + 1}",
                clause_ref=f"Table {table_idx + 1}",
                doc_name=doc_name,
                chunk_id=f"{doc_name}_table_{seq}",
            ))
            seq += 1

    return chunks


# ──────────────────────────────────────────────────────────────────────────────
# Main entry point
# ──────────────────────────────────────────────────────────────────────────────

def parse_document(path: str) -> List[Chunk]:
    """
    Auto-detect file type and return list of semantic Chunk objects.
    """
    path = str(path)
    ext = Path(path).suffix.lower()

    if ext == ".pdf":
        return parse_pdf(path)
    elif ext in (".docx", ".doc"):
        text_chunks = parse_docx(path)
        table_chunks = parse_docx_tables(path)
        return text_chunks + table_chunks
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx")
